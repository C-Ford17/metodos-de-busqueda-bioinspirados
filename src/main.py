# Bueno, la idea es pegar las imágenes en el
# word directamente desde el código, utilizando 
# unas etiquetas como [ant_colony_1]

from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from io import BytesIO
from numpy import arange
from multiprocessing import cpu_count

import time
from met_simulated_annealing import plot_learning
from met_christofides import tsp as christofides_tsp
from parallel_simulated_annealing import parallel_sa
from parallel_ant_colony import parallel_aco
from visualize_tsp import plot_tsp
from readfile import readfile

THREADS = cpu_count()
EXECUTIONS = 30

cached = []

def fill_document(template_path, output_path, data: list[dict]):
    doc = Document(template_path)

    results = {
        'sa': [],
        'aco': []
    }
    
    for i, params in enumerate(data):
        if i in cached:
            continue
        
        try:
            coords = readfile('./res/'+ params['[file]'])
        except Exception:
            coords = readfile('../res/'+ params['[file]'])

        (
            params['[sa_min_dist]'],
            params['[sa_max_dist]'],
            params['[sa_avg_dist]'],
            params['[sa_var_dist]'],
            params['[sa_avg_time]'],
            params['[sa_avg_iter]'],
            sa_min_path,
            sa_fitness_list,
            distances,
            durations
        ) = parallel_sa(coords, params, EXECUTIONS, THREADS)

        results['sa'].append((distances, durations))

        (
            params['[aco_min_dist]'],
            params['[aco_max_dist]'],
            params['[aco_avg_dist]'],
            params['[aco_var_dist]'],
            params['[aco_avg_time]'],
            params['[aco_avg_iter]'],
            aco_min_path,
            aco_fitness_list,
            distances,
            durations
        ) = parallel_aco(coords, params, EXECUTIONS, THREADS)

        results['aco'].append((distances, durations))

        for paragraph in doc.paragraphs:
            for key, value in params.items():
                paragraph.text = paragraph.text.replace(key, str(value))
    
        try:
            table = doc.tables[i]
        except Exception:
            doc.save(output_path)
            return
        for row in table.rows:
            for cell in row.cells:
                for key, value in params.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(round(value, 3)))
                for p in cell.paragraphs:
                    if '[IMAGES_SA]' in p.text:
                        p.text = p.text.replace('[IMAGES_SA]\n', '')
                        run = p.add_run()

                        buf = BytesIO()
                        plot_tsp([sa_min_path], 'Ruta minima encontrada', coords, file=buf)
                        buf.seek(0)
                        
                        run.add_picture(buf, width=Inches(3.0))

                        run = p.add_run()
                        buf = BytesIO()
                        plot_learning(sa_fitness_list, 'Aprendizaje', buf)
                        buf.seek(0)
                        run.add_picture(buf, width=Inches(3.0))
                    if '[IMAGES_ACO]' in p.text:
                        p.text = p.text.replace('[IMAGES_ACO]\n', '')
                        run = p.add_run()

                        buf = BytesIO()
                        plot_tsp([aco_min_path], 'Ruta minima encontrada', coords, file=buf)
                        buf.seek(0)
                        
                        run.add_picture(buf, width=Inches(3.0))
                        
                        run = p.add_run()
                        buf = BytesIO()
                        plot_learning(aco_fitness_list, 'Aprendizaje', buf)
                        buf.seek(0)
                        run.add_picture(buf, width=Inches(3.0))

    labels = ['SA', 'ACO']
    distances = [results['sa'][0][0], results['aco'][0][0]]
    durations = [results['sa'][0][1], results['aco'][0][1]]
    colors = ['peachpuff', 'orange']

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if "[grafico_cajas]" in run.text:
                run.text = run.text.replace("[grafico_cajas]\n", '')
                plt.clf()
                fig, ax = plt.subplots()
                ax.set_ylabel('distancia total')
                bplot = ax.boxplot(distances, patch_artist=True, tick_labels=labels)
                for patch, color in zip(bplot['boxes'], colors):
                    patch.set_facecolor(color)
                buf = BytesIO()
                plt.savefig(buf, format='png')
                buf.seek(0)
                run.add_picture(buf, width=Inches(4))
            if "[grafico_lineas]" in run.text:
                run.text = run.text.replace('[grafico_lineas]\n', '')
                plt.clf()
                plt.ylabel('distacia total')
                [a, b] = plt.plot(range(0, 30), distances[0], 'r', distances[1], 'b',)
                plt.legend([a, b], ['SA', 'ACO'])
                buf = BytesIO()
                plt.savefig(buf, format='png')
                buf.seek(0)
                run.add_picture(buf, width=Inches(8))
            if "[grafico_tiempos]" in run.text:
                run.text = run.text.replace('[grafico_tiempos]\n', '')
                plt.clf()
                X_axis = arange(1, 31)
                plt.bar(X_axis + 0.2, durations[0], 0.4, label=labels[0])
                plt.bar(X_axis - 0.2, durations[1], 0.4, label=labels[1])
                plt.xticks(arange(1, 31, 2), [str(x) for x in range(1, 31, 2)])
                plt.xlabel('ejecución')
                plt.ylabel('tiempo')
                plt.title('Duraciones')
                buf = BytesIO()
                plt.savefig(buf, format='png')
                buf.seek(0)
                run.add_picture(buf, width=Inches(8))
            

    doc.save(output_path)

    inicio = time.time()
    christofides_tsp(coords)
    fin = time.time()
    print(f"Tiempo de ejecución christofides: {fin - inicio} segundos")

if __name__ == '__main__':

    data = [
        {
            '[file]': 'uscap50.txt',
            '[temperatura_inicial]': 50**0.5,
            '[temperatura_final]': 1e-8,
            '[tasa_enfriamiento]': 0.95,
            '[sa_max_iter]': 10000,

            '[n_ants]': 10,
            '[aco_max_iter]': 20,
            '[ro]': 0.5,
            '[alpha]': 1,
            '[betha]': 1,
        },
        {
            '[file]': 'uscap50.txt',
            '[temperatura_inicial]': 1,
            '[temperatura_final]': 1e-7,
            '[tasa_enfriamiento]': 0.998,
            '[sa_max_iter]': 10000,

            '[n_ants]': 30,
            '[aco_max_iter]': 200,
            '[ro]': 0.9,
            '[alpha]': 1,
            '[betha]': 1,
        },
        {
            '[file]': 'uscap50.txt',
            '[temperatura_inicial]': 2,
            '[temperatura_final]': 1e-7,
            '[tasa_enfriamiento]': 0.9998,
            '[sa_max_iter]': 10000,

            '[n_ants]': 50,
            '[aco_max_iter]': 600,
            '[ro]': 0.8,
            '[alpha]': 1,
            '[betha]': 2,
        },
        {
            '[file]': 'uscap50.txt',
            '[temperatura_inicial]': 50**0.5,
            '[temperatura_final]': 1e-8,
            '[tasa_enfriamiento]': 0.99995,
            '[sa_max_iter]': 10000,

            '[n_ants]': 20,
            '[aco_max_iter]': 100,
            '[ro]': 0.7,
            '[alpha]': 1,
            '[betha]': 1,
        },
        {
            '[file]': 'uscap50.txt',
            '[temperatura_inicial]':50**0.5,
            '[temperatura_final]': 1e-9,
            '[tasa_enfriamiento]': 0.999995,
            '[sa_max_iter]': 10000,

            '[n_ants]': 30,
            '[aco_max_iter]': 400,
            '[ro]': 0.9,
            '[alpha]': 1,
            '[betha]': 1,
        },
        {
            '[file]': 'uscap50.txt',
            '[temperatura_inicial]':50**0.5,
            '[temperatura_final]': 1e-9,
            '[tasa_enfriamiento]': 0.999995,
            '[sa_max_iter]': 30000,

            '[n_ants]': 30,
            '[aco_max_iter]': 600,
            '[ro]': 0.9,
            '[alpha]': 1,
            '[betha]': 1,
        },
        
    ]

    try:
        fill_document('./res/template.docx', './res/output.docx', data)
        print('Se guardó el resultado en "./res/output.docx"')
    except Exception:
        fill_document('../res/template.docx', '../res/output.docx', data)
        print('Se guardó el resultado en "../res/output.docx"')

        